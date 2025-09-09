from docx import Document
from docx.shared import Cm
from PIL import Image
from io import BytesIO
import json
import sys
import os
import base64
import platform
import subprocess
import docx2pdf
import pymupdf as fitz
import random

sys.stdout.reconfigure(encoding="utf-8")


def add_photo_to_document(doc, photo_data, cell):
    try:
        image_binary = base64.b64decode(photo_data.split(",")[1])
    except:
        return

    image_stream = BytesIO(image_binary)
    image = Image.open(image_stream)

    max_width_cm, max_height_cm, dpi = 18, 13.5, 96
    max_width_px, max_height_px = int((max_width_cm / 2.54) * dpi), int((max_height_cm / 2.54) * dpi)

    aspect_ratio = image.width / image.height
    if aspect_ratio > 1:
        width_px = max_width_px
        height_px = int(width_px / aspect_ratio)
    else:
        height_px = max_height_px
        width_px = int(height_px * aspect_ratio)

    image.thumbnail((width_px, height_px))

    temp_image_path = f"temp_{os.getpid()}.jpg"
    image.save(temp_image_path)

    paragraph = cell.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(temp_image_path, width=Cm(width_px * 2.54 / dpi), height=Cm(height_px * 2.54 / dpi))

    os.remove(temp_image_path)


def generate_document(json_file):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    uploads_dir = os.path.join(script_dir, "..", "uploads", "reports")
    previews_dir = os.path.join(script_dir, "..", "uploads", "previews")
    os.makedirs(uploads_dir, exist_ok=True)
    os.makedirs(previews_dir, exist_ok=True)
    template_path = os.path.join(script_dir, "template.docx")

    with open(json_file, "r", encoding="utf-8") as f:
        user_info = json.load(f)

    doc = Document(template_path)

    if isinstance(user_info.get("works", ""), list):
        user_info["works"] = "\n• " + "\n• ".join(user_info["works"])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[дата]" in cell.text:
                    cell.text = cell.text.replace("[дата]", user_info["date"])
                if "[адрес]" in cell.text:
                    cell.text = cell.text.replace("[адрес]", user_info["address"])
                if "[назв_обор]" in cell.text:
                    cell.text = cell.text.replace("[назв_обор]", user_info["machine_name"])
                if "[номер_обор]" in cell.text:
                    cell.text = cell.text.replace("[номер_обор]", user_info["machine_number"])
                if "[инв_номер]" in cell.text:
                    cell.text = cell.text.replace("[инв_номер]", user_info["inventory_number"])
                if "[классификация]" in cell.text:
                    classification = user_info["classification"]
                    if classification == "АВ":
                        classification = "Аварийный вызов"
                    cell.text = cell.text.replace("[классификация]", classification)
                if "[материалы]" in cell.text:
                    cell.text = cell.text.replace(
                        "[материалы]", user_info.get("material", "")
                    )
                if "[рекомендации]" in cell.text:
                    cell.text = cell.text.replace(
                        "[рекомендации]", user_info.get("recommendations", "")
                    )
                if "[дефекты]" in cell.text:
                    cell.text = cell.text.replace(
                        "[дефекты]", user_info.get("defects", "")
                    )
                if "[доп_работы]" in cell.text:
                    cell.text = cell.text.replace(
                        "[доп_работы]", user_info.get("additionalWorks", "")
                    )
                if "[комментарии]" in cell.text:
                    cell.text = cell.text.replace(
                        "[комментарии]", user_info.get("comments", "")
                    )
                if "[работы]" in cell.text:
                    checklist_text = "\n".join(
                        f"• {item['task']}"
                        for item in user_info.get("checklistItems", [])
                        if item.get("done")
                    )
                    cell.text = cell.text.replace(
                        "[работы]", checklist_text if checklist_text else ""
                    )
                if "[фио]" in cell.text:
                    full_name = f"{user_info['lastName']} {user_info['firstName']}"
                    cell.text = cell.text.replace("[фио]", full_name)
                if "[вставка]" in cell.text:
                    cell.text = ""
                    for photo in user_info.get("photos", []):
                        add_photo_to_document(doc, photo, cell)
    base = f"Акт выполненных работ {user_info['date'].replace(':', '.')} {user_info['address']}"
    counter = 0
    while True:
        suffix = f" ({counter})" if counter else ""
        name = base + suffix
        docx_path = os.path.join(uploads_dir, name + ".docx")
        pdf_path = os.path.join(uploads_dir, name + ".pdf")
        preview_png_path = os.path.join(previews_dir, name + ".png")
        if not os.path.exists(docx_path) and not os.path.exists(pdf_path) and not os.path.exists(preview_png_path):
            break
        counter += 1

    doc.save(docx_path)

    pdf_converted = convert_to_pdf(docx_path)
    if not pdf_converted:
        print("Ошибка при конвертации в PDF", file=sys.stderr)
        sys.exit(1)
    
    os.replace(pdf_converted, pdf_path)
        
    final_pdf = add_stamp_to_pdf(pdf_path)
    if not final_pdf:
        print("Ошибка при добавлении печати", file=sys.stderr)
        sys.exit(1)

    preview_png = generate_preview_png(final_pdf, preview_png_path)

    filename = os.path.basename(final_pdf)
    sys.stderr.write("Generated PDF successfully\n")
    sys.stdout.write(filename + "\n")
    if preview_png:
        sys.stdout.write(os.path.basename(preview_png) + "\n")
    sys.stdout.flush()


def convert_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    system = platform.system()

    if system == "Windows":
        try:
            from docx2pdf import convert
            original_stdout = sys.stdout
            sys.stdout = sys.stderr
            convert(docx_path, pdf_path)
            sys.stdout = original_stdout
            os.remove(docx_path)
            if os.path.exists(pdf_path):
                return pdf_path
            else:
                print("Файл PDF не создан", file=sys.stderr)
                return None
        except Exception as e:
            print("Ошибка конвертации docx2pdf:", e, file=sys.stderr)
            return None
    elif system == "Linux":
        try:
            subprocess.run(
                ["unoconv", "-f", "pdf", docx_path],
                check=True,
                stderr=subprocess.PIPE
            )
            os.remove(docx_path)
            if os.path.exists(pdf_path):
                return pdf_path
            else:
                print("Файл PDF не создан", file=sys.stderr)
                return None
        except subprocess.CalledProcessError as e:
            print(f"Ошибка при конвертации через unoconv: {e}", file=sys.stderr)
            return None
    else:
        print("Неподдерживаемая ОС", file=sys.stderr)
        return None


def add_stamp_to_pdf(pdf_path):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    stamp_path = os.path.join(script_dir, "stamp.png")
    output_pdf = pdf_path.replace(".pdf", "_stamped.pdf")

    if not os.path.exists(stamp_path):
        print("Файл stamp.png не найден", file=sys.stderr)
        return None

    doc = fitz.open(pdf_path)
    
    first_page = doc[0]
    image_rect1 = fitz.Rect(100, 10, 300, 210)
    first_page.insert_image(image_rect1, filename=stamp_path, overlay=True)
    
    doc.save(output_pdf)
    doc.close()
    
    os.remove(pdf_path)
    final_pdf = pdf_path
    os.rename(output_pdf, final_pdf)
    
    return final_pdf


def generate_preview_png(pdf_path, preview_png_path):
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)
        matrix = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        tmp_path = preview_png_path + ".tmp"
        pix.save(tmp_path)
        doc.close()
        os.replace(tmp_path, preview_png_path)
        return preview_png_path
    except Exception as e:
        print(f"Ошибка генерации превью PNG: {e}", file=sys.stderr)
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Ошибка: Не передан путь к JSON-файлу", file=sys.stderr)
        sys.exit(1)

    json_file_path = sys.argv[1]
    if not os.path.exists(json_file_path):
        print("Ошибка: JSON-файл не найден", file=sys.stderr)
        sys.exit(1)

    generate_document(json_file_path)
