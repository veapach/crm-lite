"""
Модуль с основной логикой генерации документов.
Может использоваться как gRPC сервисом, так и напрямую.
"""
from docx import Document
from docx.shared import Cm
from PIL import Image
from io import BytesIO
import os
import base64
import platform
import subprocess
import pymupdf as fitz
import random


def add_photo_to_document(doc, photo_data, cell):
    """Добавление фото в документ"""
    try:
        image_binary = base64.b64decode(photo_data.split(",")[1])
    except:
        return

    image_stream = BytesIO(image_binary)
    image = Image.open(image_stream)

    # Конвертируем RGBA в RGB если нужно (JPEG не поддерживает прозрачность)
    if image.mode in ('RGBA', 'LA', 'P'):
        # Создаём белый фон
        rgb_image = Image.new('RGB', image.size, (255, 255, 255))
        if image.mode == 'P':
            image = image.convert('RGBA')
        rgb_image.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
        image = rgb_image

    max_width_cm, max_height_cm, dpi = 18, 13.5, 96
    max_width_px, max_height_px = int((max_width_cm / 2.54) * dpi), int((max_height_cm / 2.54) * dpi)

    aspect_ratio = image.width / image.height
    if aspect_ratio > 1:
        width_px = max_width_px
        height_px = int(width_px / aspect_ratio)
    else:
        height_px = max_height_px
        width_px = int(height_px * aspect_ratio)

    image.thumbnail((width_px, height_px))

    temp_image_path = f"temp_{os.getpid()}_{random.randint(1000, 9999)}.jpg"
    image.save(temp_image_path)

    paragraph = cell.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(temp_image_path, width=Cm(width_px * 2.54 / dpi), height=Cm(height_px * 2.54 / dpi))

    os.remove(temp_image_path)


def convert_to_pdf(docx_path):
    """Конвертация DOCX в PDF"""
    pdf_path = docx_path.replace(".docx", ".pdf")
    system = platform.system()

    if system == "Windows":
        try:
            from docx2pdf import convert
            import sys
            original_stdout = sys.stdout
            sys.stdout = sys.stderr
            convert(docx_path, pdf_path)
            sys.stdout = original_stdout
            os.remove(docx_path)
            if os.path.exists(pdf_path):
                return pdf_path
            else:
                return None
        except Exception as e:
            print(f"Ошибка конвертации docx2pdf: {e}")
            return None
    elif system == "Linux" or system == "Darwin":
        try:
            # Используем soffice (LibreOffice) напрямую для обеих систем
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", 
                 os.path.dirname(docx_path), docx_path],
                check=True,
                stderr=subprocess.PIPE,
                stdout=subprocess.PIPE
            )
            os.remove(docx_path)
            if os.path.exists(pdf_path):
                return pdf_path
            else:
                return None
        except subprocess.CalledProcessError as e:
            print(f"Ошибка при конвертации: {e}")
            return None
    else:
        return None


def add_stamp_to_pdf(pdf_path):
    """Добавление печати на PDF"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    stamp_path = os.path.join(script_dir, "stamp.png")
    output_pdf = pdf_path.replace(".pdf", "_stamped.pdf")

    if not os.path.exists(stamp_path):
        print("Файл stamp.png не найден")
        return None

    doc = fitz.open(pdf_path)
    
    first_page = doc[0]
    image_rect1 = fitz.Rect(100, 10, 300, 210)
    first_page.insert_image(image_rect1, filename=stamp_path, overlay=True)
    
    doc.save(output_pdf)
    doc.close()
    
    os.remove(pdf_path)
    final_pdf = pdf_path
    os.rename(output_pdf, final_pdf)
    
    return final_pdf


def generate_preview_png(pdf_path, preview_png_path):
    """Генерация PNG превью первой страницы PDF"""
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)
        # Уменьшаем масштаб с 2.0 до 1.5 для меньшего размера файла
        matrix = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        pix.save(preview_png_path)
        doc.close()
        return preview_png_path
    except Exception as e:
        print(f"Ошибка генерации превью PNG: {e}")
        return None


def generate_document_from_data(user_info: dict) -> dict:
    """
    Генерация документа на основе словаря с данными.
    Возвращает словарь с результатом:
    {
        "success": bool,
        "pdf_filename": str,
        "preview_filename": str,
        "pdf_content": bytes,
        "preview_content": bytes,
        "error": str (только при ошибке)
    }
    """
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        uploads_dir = os.path.join(script_dir, "..", "uploads", "reports")
        previews_dir = os.path.join(script_dir, "..", "uploads", "previews")
        os.makedirs(uploads_dir, exist_ok=True)
        os.makedirs(previews_dir, exist_ok=True)
        template_path = os.path.join(script_dir, "template.docx")

        if not os.path.exists(template_path):
            return {"success": False, "error": f"Шаблон не найден: {template_path}"}

        doc = Document(template_path)

        # Преобразуем works если это список
        if isinstance(user_info.get("works", ""), list):
            user_info["works"] = "\n• " + "\n• ".join(user_info["works"])

        # Заполняем шаблон
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "[дата]" in cell.text:
                        cell.text = cell.text.replace("[дата]", user_info.get("date", ""))
                    if "[адрес]" in cell.text:
                        cell.text = cell.text.replace("[адрес]", user_info.get("address", ""))
                    if "[назв_обор]" in cell.text:
                        cell.text = cell.text.replace("[назв_обор]", user_info.get("machine_name", ""))
                    if "[номер_обор]" in cell.text:
                        cell.text = cell.text.replace("[номер_обор]", user_info.get("machine_number", ""))
                    if "[инв_номер]" in cell.text:
                        cell.text = cell.text.replace("[инв_номер]", user_info.get("inventory_number", ""))
                    if "[классификация]" in cell.text:
                        classification = user_info.get("classification", "")
                        if classification == "АВ":
                            classification = "Аварийный вызов"
                        cell.text = cell.text.replace("[классификация]", classification)
                    if "[материалы]" in cell.text:
                        cell.text = cell.text.replace("[материалы]", user_info.get("material", ""))
                    if "[рекомендации]" in cell.text:
                        cell.text = cell.text.replace("[рекомендации]", user_info.get("recommendations", ""))
                    if "[дефекты]" in cell.text:
                        cell.text = cell.text.replace("[дефекты]", user_info.get("defects", ""))
                    if "[доп_работы]" in cell.text:
                        cell.text = cell.text.replace("[доп_работы]", user_info.get("additionalWorks", ""))
                    if "[комментарии]" in cell.text:
                        cell.text = cell.text.replace("[комментарии]", user_info.get("comments", ""))
                    if "[работы]" in cell.text:
                        checklist_text = "\n".join(
                            f"• {item['task']}"
                            for item in user_info.get("checklistItems", [])
                            if item.get("done")
                        )
                        cell.text = cell.text.replace("[работы]", checklist_text if checklist_text else "")
                    if "[фио]" in cell.text:
                        full_name = f"{user_info.get('lastName', '')} {user_info.get('firstName', '')}"
                        cell.text = cell.text.replace("[фио]", full_name)
                    if "[вставка]" in cell.text:
                        cell.text = ""
                        for photo in user_info.get("photos", []):
                            add_photo_to_document(doc, photo, cell)

        # Генерируем уникальное имя файла
        date_str = user_info.get("date", "").replace(":", ".")
        address_str = user_info.get("address", "")
        base = f"Акт выполненных работ {date_str} {address_str}"
        counter = 0
        while True:
            suffix = f" ({counter})" if counter else ""
            name = base + suffix
            docx_path = os.path.join(uploads_dir, name + ".docx")
            pdf_path = os.path.join(uploads_dir, name + ".pdf")
            preview_png_path = os.path.join(previews_dir, name + ".png")
            if not os.path.exists(docx_path) and not os.path.exists(pdf_path) and not os.path.exists(preview_png_path):
                break
            counter += 1

        doc.save(docx_path)

        # Конвертируем в PDF
        pdf_converted = convert_to_pdf(docx_path)
        if not pdf_converted:
            return {"success": False, "error": "Ошибка при конвертации в PDF"}
        
        os.replace(pdf_converted, pdf_path)
            
        # Добавляем печать
        final_pdf = add_stamp_to_pdf(pdf_path)
        if not final_pdf:
            return {"success": False, "error": "Ошибка при добавлении печати"}

        # Генерируем превью
        preview_png = generate_preview_png(final_pdf, preview_png_path)

        # Читаем содержимое файлов
        pdf_content = b""
        preview_content = b""
        
        if os.path.exists(final_pdf):
            with open(final_pdf, "rb") as f:
                pdf_content = f.read()
        
        if preview_png and os.path.exists(preview_png):
            with open(preview_png, "rb") as f:
                preview_content = f.read()

        return {
            "success": True,
            "pdf_filename": os.path.basename(final_pdf),
            "preview_filename": os.path.basename(preview_png) if preview_png else "",
            "pdf_content": pdf_content,
            "preview_content": preview_content,
        }

    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": f"{str(e)}\n{traceback.format_exc()}"
        }


# Для обратной совместимости - старый интерфейс через JSON файл
def generate_document(json_file):
    """Генерация документа из JSON файла (для обратной совместимости)"""
    import json
    import sys
    
    with open(json_file, "r", encoding="utf-8") as f:
        user_info = json.load(f)
    
    result = generate_document_from_data(user_info)
    
    if result["success"]:
        sys.stderr.write("Generated PDF successfully\n")
        sys.stdout.write(result["pdf_filename"] + "\n")
        if result.get("preview_filename"):
            sys.stdout.write(result["preview_filename"] + "\n")
        sys.stdout.flush()
    else:
        sys.stderr.write(result.get("error", "Unknown error") + "\n")
        sys.exit(1)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Ошибка: Не передан путь к JSON-файлу", file=sys.stderr)
        sys.exit(1)

    json_file_path = sys.argv[1]
    if not os.path.exists(json_file_path):
        print("Ошибка: JSON-файл не найден", file=sys.stderr)
        sys.exit(1)

    generate_document(json_file_path)
